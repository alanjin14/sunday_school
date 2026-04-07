"""Generate DOCX for: ECC Sunday School - Conquering Through Christ (Middle School)"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import re

# Read the markdown source
md_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "ECC Sunday School - Conquering Through Christ (Middle School).md")
with open(md_path, "r", encoding="utf-8") as f:
    md = f.read()

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

style_h1 = doc.styles['Heading 1']
style_h1.font.name = 'Calibri'
style_h1.font.size = Pt(22)
style_h1.font.bold = True
style_h1.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

style_h2 = doc.styles['Heading 2']
style_h2.font.name = 'Calibri'
style_h2.font.size = Pt(16)
style_h2.font.bold = True
style_h2.font.color.rgb = RGBColor(0x0C, 0x4A, 0x6E)

style_h3 = doc.styles['Heading 3']
style_h3.font.name = 'Calibri'
style_h3.font.size = Pt(13)
style_h3.font.bold = True
style_h3.font.color.rgb = RGBColor(0x4A, 0x19, 0x42)


def add_para(text, style='Normal', bold=False, italic=False, size=None, color=None, align=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_rich_para(parts, style='Normal', align=None):
    """Add paragraph with mixed formatting. parts = list of (text, bold, italic, color) tuples."""
    p = doc.add_paragraph(style=style)
    for text, bold, italic, color in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


# -- Title --
doc.add_heading("Sunday School: Conquering Through Christ", level=1)
add_para("ECC Redmond | April 5, 2026 Sermon Review (Easter Sunday)", italic=True,
         size=12, color=RGBColor(0x64, 0x64, 0x64))
add_para("For Middle Schoolers (Grades 6-8) | 1 Hour Session", italic=True,
         size=12, color=RGBColor(0x64, 0x64, 0x64))

# -- Leader's Overview --
doc.add_heading("LEADER'S OVERVIEW", level=2)
add_rich_para([
    ("Main Scripture: ", True, False, None),
    ("Romans 8:31-39", False, False, None),
])
add_rich_para([
    ("Speaker: ", True, False, None),
    ("Pastor Steven MacDonald", False, False, None),
])
add_rich_para([
    ("Theme: ", True, False, None),
    ("Because God paid the ultimate price for us through Christ\u2019s death and resurrection, absolutely nothing can separate us from His love \u2014 and through that love, we overwhelmingly conquer every hardship we face.", False, False, None),
])
add_rich_para([
    ("Goal: ", True, False, None),
    ("Students understand that God\u2019s love is not a feeling that comes and goes but an unstoppable power that holds them through every struggle \u2014 and that \u201cconquering\u201d doesn\u2019t mean life gets easy, but that nothing can pull them away from God.", False, False, None),
])

add_para("")
add_para("This lesson intentionally builds on last week\u2019s sermon (Elder Elton\u2019s message on becoming a living sacrifice in Romans 12), connecting back to the foundation of Romans 8 and what makes the surrendered life possible: God\u2019s unbreakable, conquering love.")
add_para("")
add_rich_para([
    ("Context: ", True, False, None),
    ("Romans 8:31-39 is the climax of Paul\u2019s argument in chapters 1-8. After covering the gospel, sin, justification, and life in the Spirit, Paul asks: \u201cWhat then shall we say to these things?\u201d He answers with five powerful rhetorical questions proving that nothing can defeat God\u2019s purpose for us, and concludes that God\u2019s love is the power that makes us more than conquerors.", False, False, None),
])

# -- Schedule Table --
doc.add_heading("SESSION SCHEDULE (60 minutes)", level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Time", "Activity", "Duration"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

schedule = [
    ("11:15", 'Icebreaker Game: "Unstoppable"', "8 min"),
    ("11:23", "Sermon Review & Teaching", "10-12 min"),
    ("11:35", "Transition to Small Groups", "2-3 min"),
    ("11:38", "Small Group Discussion", "30 min"),
    ("12:08", "Wrap-Up & Prayer", "2 min"),
]
for row_i, (time, activity, dur) in enumerate(schedule, 1):
    table.rows[row_i].cells[0].text = time
    table.rows[row_i].cells[1].text = activity
    table.rows[row_i].cells[2].text = dur
    if row_i == 4:  # bold the small group row
        for cell in table.rows[row_i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

# -- Icebreaker --
doc.add_heading('ICEBREAKER GAME: "UNSTOPPABLE" (8 min)', level=2)
doc.add_heading("How to Play", level=3)
add_para('Tell students: "We\u2019re going to play a game called UNSTOPPABLE. I\u2019m going to describe different obstacles and challenges. For each one, you have to shout out what could STOP it \u2014 or say \u2018UNSTOPPABLE\u2019 if nothing can."', italic=True)
add_para("")
add_para("Read each scenario one at a time. Let students debate and shout answers:")
add_para("")

obstacles = [
    ("A speeding train", "brakes, a wall, running out of track..."),
    ("A flood of water", "a dam, freezing it, draining it..."),
    ("A viral TikTok trend", "a new trend, getting banned, Wi-Fi going out..."),
    ("Your alarm clock on Monday morning", "the snooze button, throwing it across the room, your mom..."),
    ("A lion chasing you", "a cage, a bigger animal, running faster \u2014 probably not though..."),
    ("God\u2019s love for you", ""),
]
for obstacle, hint in obstacles:
    if hint:
        add_rich_para([
            ("\u2022  ", False, False, None),
            (obstacle, True, False, None),
            (" \u2014 What can stop it? ", False, False, None),
            ("({})".format(hint), False, True, RGBColor(0x64, 0x64, 0x64)),
        ])
    else:
        add_rich_para([
            ("\u2022  ", False, False, None),
            (obstacle, True, False, None),
            (" \u2014 What can stop it?", False, False, None),
        ])

add_para("")
add_para('Pause after the last one. Let them think. Then say: "That\u2019s exactly what Paul asked in Romans 8. He threw out every possible thing \u2014 suffering, danger, even DEATH \u2014 and the answer is: NOTHING. Absolutely nothing can separate us from God\u2019s love. That\u2019s what Pastor Steven preached about on Easter Sunday. Let\u2019s dig in."', italic=True)

# -- Sermon Review --
doc.add_heading("SERMON REVIEW & TEACHING (10-12 min)", level=2)
add_para("The goal here is not to reteach the entire sermon, but to help students remember the key ideas and prepare them for discussion.", italic=True)
add_para("")

doc.add_heading("Read Together", level=3)
add_para("Romans 8:31, 35, 37-39", bold=True)
add_para("")
p = doc.add_paragraph(style='Normal')
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run('"What then shall we say to these things? If God is for us, who is against us? ... Who will separate us from the love of Christ? Will tribulation, or trouble, or persecution, or famine, or nakedness, or danger, or sword? ... But in all these things we overwhelmingly conquer through Him who loved us. For I am convinced that neither death, nor life, nor angels, nor principalities, nor things present, nor things to come, nor powers, nor height, nor depth, nor any other created thing will be able to separate us from the love of God that is in Christ Jesus our Lord."')
run.italic = True
run.font.size = Pt(11)

add_para("")
add_para("Ask students: Paul asks five big questions here and the answer to every single one is basically \u201cNOBODY and NOTHING.\u201d Why do you think he asks them as questions instead of just saying \u201cNothing can beat God\u201d?", italic=True)

# -- Key Idea 1 --
doc.add_heading("Key Idea 1 \u2014 God Already Paid the Highest Price", level=2)
add_para("Paul\u2019s second question makes an argument from the \u201cgreater to the lesser\u201d: If God already did the HARDEST thing \u2014 giving up His own Son to die for us \u2014 why would He skip the EASY stuff?")
add_para("")
add_para('Pastor Steven used this illustration: A man who buys a half-million-dollar boat is NOT going to park it in his driveway and forget about it. He\u2019s going to fuel it, maintain it, repair it, and use it. God paid an infinitely higher price for YOU.', italic=True)
add_para("")
add_para('So when we go through hard things and wonder "Has God forgotten me?" or "Is God angry with me?" \u2014 the answer is absolutely not. He already gave everything to get you. He\u2019s not going to abandon you now.')

# -- Key Idea 2 --
doc.add_heading("Key Idea 2 \u2014 Nothing Can Condemn You or Charge You", level=2)
add_para('Pastor Steven made an important distinction: "not guilty" is not the same as "innocent."')
add_para("")
add_para("Think of it like this: if you owed someone a million dollars and couldn\u2019t pay it back, and then a friend paid the entire debt for you \u2014 you\u2019re not innocent of failing to pay. But the debt is GONE. Someone else covered it. That\u2019s what Jesus did for us. The debt of sin is fully paid. It is finished.")
add_para("")
add_para("So what do we do when we mess up? Satan wants us to wallow in guilt \u2014 to stay stuck feeling terrible. But God wants us to confess, repent, get up, and keep walking. And right now, Jesus is at the right hand of God praying for you \u2014 just like he prayed for Peter: \u201cI have prayed for you, that your faith will not fail\u201d (Luke 22:32).")

# -- Key Idea 3 --
doc.add_heading('Key Idea 3 \u2014 We Are "Super Conquerors" Through Love', level=2)
add_para('The word Paul uses in verse 37 doesn\u2019t just mean "conquer" \u2014 it means overwhelmingly conquer. Pastor Steven called this being a "super conqueror." A super conqueror doesn\u2019t just survive \u2014 he takes his enemies and makes them serve him.')
add_para("")
add_para("That\u2019s what Romans 8:28 says: \u201cAll things work together for good.\u201d God takes the very things Satan uses to try to pull us away from God and turns them into tools that draw us CLOSER to God:")
add_para("")
add_para("\u2022  Death frees us from these broken bodies and brings us to Jesus")
add_para("\u2022  Life gives us room to grow")
add_para("\u2022  Hardships teach us that this world can\u2019t satisfy \u2014 only Christ can")
add_para("\u2022  Opposition strengthens our faith")
add_para("")
add_para('We don\u2019t conquer by our own strength. We conquer through Him who loved us. "Greater is He who is in you than he who is in the world" (1 John 4:4).')
add_para("")
add_para("Pastor Steven closed with the story of C.H. Spurgeon \u2014 one of the greatest preachers in history, who suffered lifelong depression and was often bedridden with gout. Yet he never gave up. He kept turning to God\u2019s Word, kept preaching, kept serving. His ministry planted over 200 churches. He conquered through Christ \u2014 not by being \u201chappy and clappy\u201d all the time, but by holding on to God\u2019s love and never letting go.")

# -- Small Group Discussion --
doc.add_heading("SMALL GROUP DISCUSSION (30 min)", level=2)
add_para("Split students into groups of 4-5 with one leader.", bold=True)
add_para("Leaders should focus on conversation rather than rushing through questions. Let the good conversations breathe \u2014 it\u2019s OK to skip some questions.")
add_para("Encourage students to share honestly and ask questions. No judgment zone.")

# Warm-Up
doc.add_heading("Warm-Up Questions", level=3)
add_para("1. What\u2019s ONE thing you remember from Pastor Steven\u2019s Easter sermon? (Even a random detail counts!)")
add_para("")
add_para("2. Let\u2019s do a quick poll \u2014 raise your hand for your answer. When you\u2019re going through a really hard time, what\u2019s the FIRST thing you usually do?")
add_para("     (A) Talk to a friend or family member")
add_para("     (B) Try to figure it out on your own")
add_para("     (C) Scroll on your phone to distract yourself")
add_para("     (D) Pray about it")
add_para("     (No judgment \u2014 be honest!) Now: what does it say about us that praying usually isn\u2019t our first move, even though God says nothing can separate us from His love?", italic=True)
add_para("")
add_para("3. Was there anything that confused you or that you want to ask about?")

# Understanding
doc.add_heading("Understanding the Topic", level=3)
add_para('4. Pastor Steven said being "not guilty" is NOT the same as being "innocent." He used the example of a servant whose massive debt was paid by his master. Why is that difference important? How does it change the way you think about your own mistakes and sins?')
add_para("")
add_para('5. The sermon said we are "super conquerors" \u2014 meaning God takes the bad things in our lives and makes them serve His purpose for our good. Can you think of a time when something hard or painful actually helped you grow or taught you something important?')
add_para("")
add_para("6. Pastor Steven told the story of C.H. Spurgeon, who suffered from depression and physical pain his whole life but never stopped following Christ. What does that tell us about what \u201cconquering\u201d actually looks like? Does it mean we have to feel strong and happy all the time?")

# Faith vs Real Life
doc.add_heading("Faith vs Real Life", level=3)
add_para("7. Paul lists things that COULD separate us from God\u2019s love: tribulation, trouble, persecution, famine, danger, sword. Let\u2019s update that list for middle school life. What are the things that make YOU feel far from God or make you question whether God cares? Think about:")
add_para("     \u2022  A really bad day at school")
add_para("     \u2022  Friends turning on you or drama in the group chat")
add_para("     \u2022  Feeling like your prayers aren\u2019t being answered")
add_para("     \u2022  Comparing yourself to others on social media")
add_para("     \u2022  Family stress at home")
add_para("     \u2022  Feeling alone even in a crowd")
add_para("")
add_para('8. Pastor Steven said that when we struggle, we sometimes wonder: "Has God forgotten me? Is God angry with me?" Be honest \u2014 have you ever thought that? What was happening when you felt that way? How does knowing that God "already paid the highest price for you" change those thoughts?')
add_para("")
add_para("9. The sermon said Satan wants us to wallow in guilt when we mess up, but God wants us to confess, get up, and keep walking. Think about the last time you felt really guilty about something. Did you get stuck in that guilt, or did you bring it to God and move forward? What\u2019s the difference between healthy conviction (that leads to repentance) and destructive guilt (that keeps you stuck)?")

# Real Life Application
doc.add_heading("Real Life Application", level=3)
add_para("10. Pastor Steven gave three ways to respond to God\u2019s conquering love:")
add_para("     (A) Bring Easter into your daily life \u2014 see difficulties as opportunities to rely on God")
add_para("     (B) Love and serve God \u2014 read His Word, follow His commands")
add_para("     (C) Care for those God cares for \u2014 share burdens, pray for people, share Christ")
add_para("     Which of these three is HARDEST for you right now? Which one feels most natural? Pick the hardest one and describe what it would look like in your life this week.", italic=True)
add_para("")
add_para('11. Pastor Steven said: "You don\u2019t have to be happy and clappy all the time to be faithful to Christ. You just have to keep following him and trusting him." Think about a current struggle in your life. What does "not letting go of God" look like in THAT specific situation? What\u2019s one practical thing you can do when you feel like giving up?')

# Personal Reflection
doc.add_heading("Personal Reflection", level=3)
add_para('12. Paul says he is "convinced" that nothing can separate us from God\u2019s love. On a scale of 1-10, how convinced are YOU? What would it take to move that number up? Is it about knowing more, or experiencing more, or something else?')
add_para("")
add_para("13. The sermon said God takes the things the devil would use to separate us from Him and instead uses them to draw us closer. If you believed Romans 8:28 was 100% true for YOUR life, how would you handle your biggest current challenge differently? Share one thing you want to hold onto from this lesson going forward.")
add_para("     (Leaders: have each student share theirs out loud.)", italic=True)

# -- Wrap-Up --
doc.add_heading("WRAP-UP & PRAYER (2 min)", level=2)
add_para("Ask students to finish this sentence out loud or in their heads:")
add_para("")
p = doc.add_paragraph(style='Normal')
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run('"Nothing can separate me from God\u2019s love \u2014 not even __________."')
run.bold = True
run.italic = True

add_para("")
add_para("Close with prayer:", bold=True)
add_para("")
add_para('"God, thank you for Easter. Thank you that you didn\u2019t just conquer death 2,000 years ago and leave us to figure out the rest on our own. Thank you that you bring Easter into our lives TODAY \u2014 that the same power that raised Jesus from the dead is alive in us through your Holy Spirit. Pastor Steven reminded us that you already paid the highest price for us, and you will never forget us, never abandon us, never stop fighting for us. When we struggle \u2014 when we feel guilty, when we feel alone, when we wonder if you even care \u2014 help us remember that NOTHING in all creation can separate us from your love. Not our worst day. Not our biggest failure. Not our deepest fear. Help us be like Spurgeon \u2014 not perfect, not always happy, but never letting go of you. Make us more than conquerors this week. Not by our own strength, but through your love. In Jesus\u2019 name, amen."', italic=True)

# -- Save --
out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "ECC Sunday School - Conquering Through Christ (Middle School).docx")
doc.save(out)
print("Saved -> {}".format(out))

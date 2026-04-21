"""Generate DOCX for: ECC Sunday School - The Power to Scatter (Middle School)"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

style_h1 = doc.styles['Heading 1']
style_h1.font.name = 'Calibri'
style_h1.font.size = Pt(22)
style_h1.font.bold = True
style_h1.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

style_h2 = doc.styles['Heading 2']
style_h2.font.name = 'Calibri'
style_h2.font.size = Pt(16)
style_h2.font.bold = True
style_h2.font.color.rgb = RGBColor(0x0C, 0x4A, 0x6E)

style_h3 = doc.styles['Heading 3']
style_h3.font.name = 'Calibri'
style_h3.font.size = Pt(13)
style_h3.font.bold = True
style_h3.font.color.rgb = RGBColor(0x4A, 0x19, 0x42)


def add_para(text, style='Normal', bold=False, italic=False, size=None, color=None, align=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_rich_para(parts, style='Normal', align=None):
    """Add paragraph with mixed formatting. parts = list of (text, bold, italic, color) tuples."""
    p = doc.add_paragraph(style=style)
    for text, bold, italic, color in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


# -- Title --
doc.add_heading("Sunday School: Scattered, Not Shattered \u2014 The Power to Scatter", level=1)
add_para("ECC Redmond | April 19, 2026 Sermon Review", italic=True,
         size=12, color=RGBColor(0x64, 0x64, 0x64))
add_para("For Middle Schoolers (Grades 6-8) | 1 Hour Session", italic=True,
         size=12, color=RGBColor(0x64, 0x64, 0x64))

# -- Leader's Overview --
doc.add_heading("LEADER\u2019S OVERVIEW", level=2)
add_rich_para([
    ("Main Scripture: ", True, False, None),
    ("Acts 8:1-8", False, False, None),
])
add_rich_para([
    ("Speaker: ", True, False, None),
    ("Pastor Steve Moy", False, False, None),
])
add_rich_para([
    ("Theme: ", True, False, None),
    ("God uses disruption and persecution not to destroy His people but to deploy them \u2014 scattering ordinary believers to spread the gospel and bring unexpected joy to places that need it most.", False, False, None),
])
add_rich_para([
    ("Goal: ", True, False, None),
    ("Students understand that God can use hard, unwanted changes in their lives to put them exactly where He wants them \u2014 and that they don\u2019t need a special title or training to share Jesus with the people around them.", False, False, None),
])

add_para("")
add_para('This lesson connects to the series "The Gospel: Power of God" in Acts, following Stephen\u2019s martyrdom (Acts 7) and showing how the church\u2019s worst crisis became the gospel\u2019s greatest expansion.')
add_para("")
add_rich_para([
    ("Context: ", True, False, None),
    ("After Stephen was stoned to death, Saul (who took pleasure in killing Christians \u2014 the Greek word means \u201cto be pleased with\u201d) launched a fierce, unprecedented persecution. He went house to house, dragging off men AND women, imprisoning and killing believers. The Christians scattered throughout Judea and Samaria. But instead of hiding, they preached wherever they went. Philip \u2014 a deacon who served tables (Acts 6), not an apostle \u2014 went to Samaria (a place Jews had despised for centuries) and proclaimed Christ. The result: healing, deliverance, and \u201cmuch joy\u201d in the city. The early Christians were scattered but not shattered.", False, False, None),
])

# -- Schedule Table --
doc.add_heading("SESSION SCHEDULE (60 minutes)", level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Time", "Activity", "Duration"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

schedule = [
    ("11:15", 'Icebreaker Game: "Scattered!"', "8 min"),
    ("11:23", "Sermon Review & Teaching", "10-12 min"),
    ("11:35", "Transition to Small Groups", "2-3 min"),
    ("11:38", "Small Group Discussion", "30 min"),
    ("12:08", "Wrap-Up & Prayer", "2 min"),
]
for row_i, (time, activity, dur) in enumerate(schedule, 1):
    table.rows[row_i].cells[0].text = time
    table.rows[row_i].cells[1].text = activity
    table.rows[row_i].cells[2].text = dur
    if row_i == 4:
        for cell in table.rows[row_i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

# -- Icebreaker --
doc.add_heading('ICEBREAKER GAME: "SCATTERED!" (8 min)', level=2)
doc.add_heading("How to Play", level=3)
add_para('Tell students: "We\u2019re going to play a game called SCATTERED! I\u2019m going to name something that starts together in one place. Your job: shout out what could SCATTER it!"', italic=True)
add_para("")
add_para("Read each scenario one at a time. Let students debate and shout answers:")
add_para("")

obstacles = [
    ("A pile of leaves on the ground", "wind, a leaf blower, jumping in them, a dog running through..."),
    ("A flock of pigeons in the park", "a loud noise, a kid running at them, a dog, a cat..."),
    ("Your friend group at school", "summer break, moving to a new school, drama, different lunch periods..."),
    ("A bag of marbles dropped on the floor", "gravity, a slippery floor, chaos..."),
    ("Seeds in a dandelion puff", "a breath, the wind, a kid blowing on it..."),
]
for obstacle, hint in obstacles:
    add_rich_para([
        ("\u2022  ", False, False, None),
        (obstacle, True, False, None),
        (" \u2014 What scatters them? ", False, False, None),
        ("({})".format(hint), False, True, RGBColor(0x64, 0x64, 0x64)),
    ])

add_para("")
add_para('Pause after the dandelion one and say: "Notice something about the dandelion \u2014 when the seeds scatter, they don\u2019t just fly away and die. They LAND somewhere and grow into NEW flowers."', italic=True)
add_para("")
add_rich_para([
    ("\u2022  ", False, False, None),
    ("The early church in Jerusalem", True, False, None),
    (" \u2014 What scattered them? ", False, False, None),
    ("(persecution!)", False, True, RGBColor(0x64, 0x64, 0x64)),
])
add_para("")
add_para('Say: "After Stephen was killed for his faith, a man named Saul started going house to house, dragging Christians out and throwing them in prison \u2014 or worse. The believers scattered \u2014 they ran for their lives. But here\u2019s the crazy part: like dandelion seeds, everywhere they landed, the gospel GREW. Pastor Steve said they were \u2018scattered but not shattered.\u2019 That\u2019s what we\u2019re looking at today."', italic=True)

# -- Sermon Review --
doc.add_heading("SERMON REVIEW & TEACHING (10-12 min)", level=2)
add_para("The goal here is not to reteach the entire sermon, but to help students remember the key ideas and prepare them for discussion.", italic=True)
add_para("")

doc.add_heading("Read Together", level=3)
add_para("Acts 8:1, 4-8", bold=True)
add_para("")
p = doc.add_paragraph(style='Normal')
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run('"And Saul approved of his execution. And there arose on that day a great persecution against the church in Jerusalem, and they were all scattered throughout the regions of Judea and Samaria, except the apostles. ... Now those who were scattered went about preaching the word. Philip went down to the city of Samaria and proclaimed to them the Christ. And the crowds with one accord paid attention to what was being said by Philip, when they heard him and saw the signs that he did. For unclean spirits, crying out with a loud voice, came out of many who had them, and many who were paralyzed or lame were healed. So there was much joy in that city."')
run.italic = True
run.font.size = Pt(11)

add_para("")
add_para("Ask students: Look at the beginning of this passage (persecution, prison, death) and the end (healing, crowds paying attention, much joy). What happened in between to get from point A to point B?", italic=True)

# -- Key Idea 1 --
doc.add_heading("Key Idea 1 \u2014 God Uses Disruption to Move His People", level=2)
add_para('Jesus told His disciples in Acts 1:8: "You will be my witnesses in Jerusalem, and in all Judea and Samaria, and to the ends of the earth." But the church had stayed in Jerusalem. They were comfortable \u2014 gathering together, eating together, worshiping together, watching the community grow. Sound familiar? It\u2019s kind of like our retreats or youth group nights \u2014 being surrounded by other Christians just feels good.')
add_para("")
add_para("But they weren\u2019t going anywhere. They had a mission from Jesus and they weren\u2019t doing it.")
add_para("")
add_para('Then persecution hit. Saul \u2014 who the Greek text says was "pleased with" Stephen\u2019s murder \u2014 started going house to house, dragging off men and women. The word "ravaging" in Greek carries two images: an army destroying a city and a wild animal tearing at its meat. That\u2019s how vicious this was.')
add_para("")
add_para("Pastor Steve used a Star Wars comparison: Luke Skywalker was minding his own business on Tatooine, harvesting water, living a comfortable life. Then two droids showed up, his family was killed, his home destroyed \u2014 and that painful disruption pushed him into something greater: the Rebellion. In the same way, the early church didn\u2019t choose to leave Jerusalem, but God used that disruption to send them exactly where Jesus had told them to go.", italic=True)
add_para("")
add_para("Sometimes the hardest disruptions in our lives \u2014 a move, a broken friendship, a school change, a family crisis \u2014 are God repositioning us to be exactly where He wants us.")

# -- Key Idea 2 --
doc.add_heading("Key Idea 2 \u2014 Ordinary People, Extraordinary Mission", level=2)
add_para('Here\u2019s the surprising detail in verse 1: "they were all scattered... except the apostles." The apostles \u2014 the trained leaders, the inner circle who followed Jesus \u2014 stayed in Jerusalem. The people who scattered and preached the word were ordinary believers. Students, teachers, cooks, regular people. Not the trained preachers.')
add_para("")
add_para('Philip is the star of this chapter. Pastor Steve had fun with this \u2014 he showed pictures of famous "Phils" (Dr. Phil, Phil Jackson the basketball coach, Uncle Phil from Fresh Prince) before revealing the real Philip. He wasn\u2019t an apostle or a theologian. He was one of the seven deacons chosen in Acts 6 to serve food to widows. Pastor Steve put it bluntly: "In essence, he was a waiter."', italic=True)
add_para("")
add_para("Yet Philip went to Samaria \u2014 a place Jews had despised and avoided for hundreds of years due to deep ethnic and religious tension going back to Ezra in the Old Testament. The woman at the well told Jesus: \u201cJews do not have dealings with Samaritans.\u201d For a Jewish Christian to go there and preach Christ took incredible courage.")
add_para("")
add_para('How did they do it? Not by their own power. 2 Timothy 1:7 says: "For God gave us a spirit not of fear but of power and of love and self-discipline." The Holy Spirit compelled them.')
add_para("")
add_para("Pastor Steve also drew a parallel to the COVID-19 pandemic (2020-2022): churches were forcibly closed, believers scattered to living rooms and Zoom screens. There was a surge in digital missions and neighborhood care. The word was taken out of the temples and into the streets \u2014 just like Acts 8. Hundreds of people still tune in to ECC\u2019s messages online because of that scattering.")
add_para("")
add_para('Tim Keller said: "The Gospel is this; We are more sinful and flawed than we ever dared to believe, yet more loved and accepted than we ever dared to hope." This gospel is not reserved for pastors and missionaries \u2014 it\u2019s entrusted to all of us.', italic=True)

# -- Key Idea 3 --
doc.add_heading("Key Idea 3 \u2014 Scattering Leads to Unexpected Joy", level=2)
add_para('Look at how the passage ends: "So there was much joy in that city" (v. 8). Pastor Steve emphasized that this joy was not private or quiet \u2014 it was overflowing, communal, and visible.')
add_para("")
add_para("The story starts with a funeral (Stephen\u2019s burial), prison, and persecution. It ends with healing, deliverance from evil spirits, and a whole city full of joy. God turned the church\u2019s worst day into the gospel\u2019s greatest advance.")
add_para("")
add_para('Pastor Steve said: "The gospel message doesn\u2019t just fix problems \u2014 it restores joy to all the communities." When God\u2019s kingdom breaks in, it\u2019s a shared, collective restoration.', italic=True)
add_para("")
add_para('He quoted Psalm 65:8: "The whole earth is filled with awe at your wonders. Where morning dawns, where evening fades, you call forth songs of joy." This joy comes from God \u2014 it\u2019s not just a feeling but an overwhelming certainty and hope.')
add_para("")
add_para('C.S. Lewis said: "Joy is the serious business of heaven." The joy in Samaria was the serious, deep, real joy of heaven breaking into a city no one expected God to reach.', italic=True)
add_para("")
add_para("The pattern is simple: Persecution \u2192 Scattering \u2192 Preaching \u2192 Healing \u2192 JOY. When God scatters us, He\u2019s planting us \u2014 and the harvest is often more joyful than anything we could have planned.")

# -- Small Group Discussion --
doc.add_heading("SMALL GROUP DISCUSSION (30 min)", level=2)
add_para("Split students into groups of 4-5 with one leader.", bold=True)
add_para("Leaders should focus on conversation rather than rushing through questions. Let the good conversations breathe \u2014 it\u2019s OK to skip some questions.")
add_para("Encourage students to share honestly and ask questions. No judgment zone.")

# Warm-Up
doc.add_heading("Warm-Up Questions", level=3)
add_para("1. What\u2019s ONE thing you remember from Pastor Steve\u2019s sermon? (Even a random detail counts \u2014 bonus points if you remember the famous Phils or the Star Wars reference!)")
add_para("")
add_para("2. Have you ever been \u201cscattered\u201d \u2014 moved to a new school, lost a friend group, had your routine completely disrupted? How did it feel at the time? Looking back, did anything good come out of it?")
add_para("")
add_para("3. Was there anything from the sermon or scripture that confused you or that you want to ask about?")

# Understanding
doc.add_heading("Understanding the Topic", level=3)
add_para("4. The believers in Jerusalem were comfortable \u2014 they had great community and the church was growing. They were gathering, eating together, worshiping. But they weren\u2019t going to Judea and Samaria like Jesus told them to. Why do you think they stayed put? What does that tell us about comfort zones and God\u2019s mission?")
add_para("")
add_para("5. The apostles (the \u201cprofessionals\u201d) stayed in Jerusalem. The ordinary believers were the ones who scattered and preached. Pastor Steve said mission \u201cbelongs to all of us\u201d \u2014 not just pastors and missionaries. Why is that such a big deal? What does it say about who God uses?")
add_para("")
add_para("6. Philip went to Samaria \u2014 a place Jews had avoided and looked down on for hundreds of years because of ethnic and religious tension. That would be like going to the group at school that everyone avoids or talks about behind their backs. Why did it take so much courage for Philip to go there? What made him willing to do it?")

# Faith vs Real Life
doc.add_heading("Faith vs Real Life", level=3)
add_para("7. Pastor Steve compared the early church\u2019s disruption to Luke Skywalker\u2019s story \u2014 comfortable life on Tatooine, then everything gets upended, and that pushes him into a bigger purpose. Have you ever gone through something hard (a move, a friendship ending, a family change) that ended up pushing you into something you wouldn\u2019t have experienced otherwise? Share if you\u2019re comfortable.")
add_para("")
add_para("8. Pastor Steve talked about the COVID pandemic as a modern-day scattering \u2014 churches closed, believers scattered to Zoom and living rooms, and the gospel actually spread MORE through digital missions. Can you think of other times when something bad or disruptive in your life actually opened a door for something good? Does that make the bad thing \u201cworth it,\u201d or is it more complicated than that?")
add_para("")
add_para('9. The application said: "See yourself on mission. Be a presence of peace, integrity, and grace." Let\u2019s be real \u2014 where are you \u201cscattered\u201d right now? Think about your school, your teams, your neighborhood, your online spaces. Do you see those places as your mission field, or just places you happen to be? What\u2019s the difference?')

# Real Life Application
doc.add_heading("Real Life Application", level=3)
add_para("10. The sermon\u2019s application gave three challenges:")
add_para("     (A) Equip everyone to share the Gospel \u2014 be bridge builders across cultures and generations")
add_para("     (B) See yourself on mission \u2014 be a presence of peace, integrity, and grace wherever you are")
add_para("     (C) Live with courage beyond your comfort zone \u2014 build relationships with people who are different from you")
add_para("     Which of these three feels most relevant to YOUR life right now? Pick one and describe what it would look like \u2014 specifically \u2014 at school, at home, or with friends this week. (Not \u201cbe nice to people\u201d \u2014 be specific: WHO, WHERE, WHEN.)", italic=True)
add_para("")
add_para('11. Think about your "Samaria" \u2014 a group, a person, or a place that you normally avoid or that feels uncomfortable. It could be the kid who sits alone at lunch, a family member you don\u2019t get along with, or a group at school that\u2019s very different from yours. Philip was "in essence a waiter" \u2014 nothing special on paper \u2014 but he crossed a barrier that had been there for hundreds of years. What would it look like to be a "Philip" this week? What would it cost you? What might God do through it?')

# Personal Reflection
doc.add_heading("Personal Reflection", level=3)
add_para('12. The early church didn\u2019t choose to be scattered \u2014 it happened to them. But they chose what to do with it (preach the word wherever they went, instead of hiding). Think about something in your life right now that you didn\u2019t choose \u2014 a situation you\u2019d rather not be in. What would it look like to "preach the word" (live out your faith and share hope) in THAT situation instead of just surviving it? How does 2 Timothy 1:7 apply here: "God gave us a spirit not of fear but of power and of love and self-discipline"?')
add_para("")
add_para('13. Acts 8:8 says "there was much joy in that city." Not just a little joy \u2014 MUCH joy. It wasn\u2019t private or quiet \u2014 it was overflowing and communal. C.S. Lewis said "Joy is the serious business of heaven." When was the last time you experienced real, deep joy \u2014 not just fun or entertainment, but the kind of joy that feels like God is in it? What was happening? How can you be someone who brings that kind of joy to the people and places where God has scattered you?')
add_para("     (Leaders: have each student share one thing they want to take away from this lesson.)", italic=True)

# -- Wrap-Up --
doc.add_heading("WRAP-UP & PRAYER (2 min)", level=2)
add_para("Ask students to finish this sentence out loud or in their heads:")
add_para("")
p = doc.add_paragraph(style='Normal')
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run('"God might be scattering me to __________ so that I can __________."')
run.bold = True
run.italic = True

add_para("")
add_para("Close with prayer:", bold=True)
add_para("")
add_para('"God, thank you for this passage in Acts. Thank you that when the early Christians were scattered, they were not shattered \u2014 you still had a mission for them. You used persecution to push your people exactly where Jesus told them to go. Help us see our own \u2018scattering\u2019 the same way \u2014 the new school, the broken friendship, the hard season, the uncomfortable situation \u2014 not as punishment but as your plan to put us where you need us. Give us the courage of Philip \u2014 a regular person, a table server, who crossed barriers that had been there for centuries and proclaimed Christ to people everyone else avoided. We don\u2019t need a title or a degree. We just need to know you and be willing to share you. You gave us a spirit not of fear but of power and of love and self-discipline. Help us be bridge builders this week. Help us step outside our comfort zones. Help us bring your joy \u2014 overflowing, visible, real joy \u2014 wherever we land. Scattered, but never shattered. In Jesus\u2019 name, amen."', italic=True)

# -- Save --
out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "ECC Sunday School - The Power to Scatter (Middle School).docx")
doc.save(out)
print("Saved -> {}".format(out))

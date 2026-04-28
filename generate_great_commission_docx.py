"""Generate DOCX for: ECC Sunday School - Your Mission in the Great Commission (Middle School)"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

style_h1 = doc.styles['Heading 1']
style_h1.font.name = 'Calibri'
style_h1.font.size = Pt(22)
style_h1.font.bold = True
style_h1.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

style_h2 = doc.styles['Heading 2']
style_h2.font.name = 'Calibri'
style_h2.font.size = Pt(16)
style_h2.font.bold = True
style_h2.font.color.rgb = RGBColor(0x0C, 0x4A, 0x6E)

style_h3 = doc.styles['Heading 3']
style_h3.font.name = 'Calibri'
style_h3.font.size = Pt(13)
style_h3.font.bold = True
style_h3.font.color.rgb = RGBColor(0x4A, 0x19, 0x42)


def add_para(text, style='Normal', bold=False, italic=False, size=None,
             color=None, align=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_rich_para(parts, style='Normal', align=None):
    """parts = list of (text, bold, italic, color) tuples."""
    p = doc.add_paragraph(style=style)
    for text, bold, italic, color in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


# -- Title --
doc.add_heading("Sunday School: Your Mission in the Great Commission", level=1)
add_para("ECC Redmond | April 26, 2026 Sermon Review", italic=True,
         size=12, color=RGBColor(0x64, 0x64, 0x64))
add_para("For Middle Schoolers (Grades 6-8) | 1 Hour Session", italic=True,
         size=12, color=RGBColor(0x64, 0x64, 0x64))

# -- Leader's Overview --
doc.add_heading("LEADER’S OVERVIEW", level=2)
add_rich_para([
    ("Main Scripture: ", True, False, None),
    ("Acts 13:1-5", False, False, None),
])
add_rich_para([
    ("Speaker: ", True, False, None),
    ("Elder Elton Lee", False, False, None),
])
add_rich_para([
    ("Theme: ", True, False, None),
    ("The Holy Spirit initiates God’s global mission and the church responds — both as senders and as the sent. Every believer is part of the Great Commission, and we discover our specific role through Learn → Discern → Confirm.", False, False, None),
])
add_rich_para([
    ("Goal: ", True, False, None),
    ('Students understand that "missions" isn’t just for adults in foreign countries — every Christian has a mission, every day. They learn that the Holy Spirit is the One who initiates and that responding starts with listening (not just praying for direction once and giving up).', False, False, None),
])

add_para("")
add_para('This lesson connects to last week’s message ("Scattered, Not Shattered" — Acts 8). The persecuted believers who scattered to Antioch became the very church that now sends Barnabas and Paul out to the Gentile world. We’re watching Acts 1:8 unfold from "Jerusalem" → "Judea and Samaria" → "the ends of the earth."')
add_para("")
add_rich_para([
    ("Context: ", True, False, None),
    ("Acts 13 opens at the Antioch church — the first majority-Gentile church (Acts 11). While the leaders were worshiping and fasting, the Holy Spirit told the church to set apart Barnabas and Saul (Paul) for missionary work. Notice the setting — it wasn’t a strategy meeting. It was during normal spiritual rhythms. The church laid hands on them, sent them off, and the Holy Spirit Himself sent them on their way to Cyprus. This launches Paul’s first missionary journey — and the rest of the book of Acts unfolds from here. The big idea: ", False, False, None),
    ("God already has a plan; we just need to figure out how to be part of it.", True, False, None),
])

# -- Schedule Table --
doc.add_heading("SESSION SCHEDULE (60 minutes)", level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Time", "Activity", "Duration"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

schedule = [
    ("11:15", 'Icebreaker Game: "Mission: Possible!"', "8 min"),
    ("11:23", "Sermon Review & Teaching", "10-12 min"),
    ("11:35", "Transition to Small Groups", "2-3 min"),
    ("11:38", "Small Group Discussion", "30 min"),
    ("12:08", "Wrap-Up & Prayer", "2 min"),
]
for row_i, (time, activity, dur) in enumerate(schedule, 1):
    table.rows[row_i].cells[0].text = time
    table.rows[row_i].cells[1].text = activity
    table.rows[row_i].cells[2].text = dur
    if row_i == 4:
        for cell in table.rows[row_i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

# -- Icebreaker --
doc.add_heading('ICEBREAKER GAME: "MISSION: POSSIBLE!" (8 min)', level=2)

doc.add_heading("Setup", level=3)
add_para("You’ll need: paper and pens, or just shouted-out answers. Optional: dramatic Mission: Impossible theme music in the background.")

doc.add_heading("How to Play", level=3)
add_para('Tell students: "This week’s sermon was about MISSIONS. Not Mission Impossible — Mission POSSIBLE. Every one of you is on a mission whether you realize it or not. Let’s figure out what some missions are."', italic=True)
add_para("")

add_para('Round 1 — "Who’s the Mission?" (3 min)', bold=True)
add_para("Read these famous \"missions\" out loud and have students shout out the character / team / agency:")
add_para("")

round1 = [
    ("\U0001f9b8", '"With great power comes great responsibility. Mission: protect New York."', "Spider-Man"),
    ("\U0001f680", '"That’s one small step for man, one giant leap for mankind. Mission: walk on the moon."', "Apollo 11 / Neil Armstrong"),
    ("\U0001f355", '"In a half-shell, in a sewer, save the city from Shredder. Mission: protect NYC."', "TMNT"),
    ("\U0001f420", '"Cross the entire ocean to find one little fish. Mission: rescue your son."', "Marlin (Finding Nemo)"),
    ("⚡", '"Take a magic ring to a volcano and throw it in. Mission: destroy the One Ring."', "Frodo / The Fellowship"),
    ("\U0001f31f", '"Go therefore and make disciples of all nations." Mission: ?', "Every Christian"),
]
for emoji, line, answer in round1:
    add_rich_para([
        ("{}  ".format(emoji), False, False, None),
        (line, False, True, None),
        ("  →  ", False, False, None),
        (answer, True, False, None),
    ])

add_para("")
add_para('Pause on that last one. Say: "Wait — really? That’s MY mission? I’m 12. I haven’t even finished pre-algebra."', italic=True)
add_para("")

add_para('Round 2 — "Mission Receivers" (3 min)', bold=True)
add_para('"In every mission story, there’s a moment where the hero gets the mission. How is each one delivered?" (Have students match — quick rapid-fire.)', italic=True)
add_para("")

round2 = [
    ("\U0001f3ac", "Mission: Impossible", "self-destructing tape"),
    ("\U0001f3f0", "Frodo", "handed to him at the Council of Elrond"),
    ("\U0001f30c", "Star Wars", "hidden inside R2-D2 as a hologram"),
    ("\U0001f4dc", "Avengers", "Nick Fury shows up at your door"),
    ("\U0001f54a️", "Acts 13", "the Holy Spirit speaks while the church is worshiping and fasting"),
]
for emoji, name, how in round2:
    add_rich_para([
        ("{}  ".format(emoji), False, False, None),
        (name, True, False, None),
        ("  →  ", False, False, None),
        (how, False, True, None),
    ])

add_para("")
add_para('Say: "Notice the difference. The hero doesn’t usually get a mission while sitting around watching TikTok. The mission shows up when they’re DOING something — paying attention. In Acts 13, the church wasn’t having a missions strategy meeting. They were just doing what they always did — worshiping and fasting — and the Holy Spirit spoke."', italic=True)

doc.add_heading("Transition Question (1-2 min)", level=3)
add_rich_para([
    ("Ask: ", False, False, None),
    ('"If God wanted to give YOU a mission this week, would you actually be paying enough attention to hear Him?"', False, True, None),
])
add_para("")
add_para('Don’t let students answer too fast. Let it sit. Then say: "That’s exactly what we’re going to talk about today."', italic=True)

# -- Sermon Review --
doc.add_heading("SERMON REVIEW & TEACHING (10-12 min)", level=2)
add_para("The goal here is not to reteach the entire sermon, but to help students remember the key ideas and prepare them for discussion.", italic=True)
add_para("")

doc.add_heading("Read Together", level=3)
add_para("Acts 13:1-5", bold=True)
add_para("")
p = doc.add_paragraph(style='Normal')
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run('"Now in the church at Antioch there were prophets and teachers: Barnabas, Simeon called Niger, Lucius of Cyrene, Manaen (who had been brought up with Herod the tetrarch) and Saul. While they were worshiping the Lord and fasting, the Holy Spirit said, ‘Set apart for me Barnabas and Saul for the work to which I have called them.’ So after they had fasted and prayed, they placed their hands on them and sent them off. The two of them, sent on their way by the Holy Spirit, went down to Seleucia and sailed from there to Cyprus. When they arrived at Salamis, they proclaimed the word of God in the Jewish synagogues. John was with them as their helper."')
run.italic = True
run.font.size = Pt(11)

add_para("")
add_para("Ask:", bold=True)
add_para("•  Who’s doing the sending in this passage? Look closely at v. 2 and v. 4. (The Holy Spirit. He speaks in v. 2 and is described as the one who sent them in v. 4.)")
add_para("•  What were the people doing when the Holy Spirit spoke? (Worshiping and fasting — not having a missions meeting.)")

# -- Key Idea 1 --
doc.add_heading("Key Idea 1 — The Holy Spirit Initiates God’s Global Mission", level=2)
add_para("Last week (in Acts 8) we saw believers scattered by persecution — they ran for their lives and ended up planting the church in Antioch. Here in Acts 13, that same Antioch church is now sending out missionaries to the rest of the world. Pretty crazy, right? The persecuted became the senders.")
add_para("")
add_para('But notice WHO is calling the shots. The Holy Spirit is the one who initiates the mission. The church didn’t sit around and brainstorm a strategy. They didn’t put up a flyer that said "Wanted: 2 missionaries, must be willing to travel." They were just worshiping and fasting — doing what they were already doing — and the Holy Spirit interrupted with, "Set apart Barnabas and Saul for the work I’ve called them to."')
add_para("")
add_para("You can see the Holy Spirit’s fingerprints all over Paul’s missionary journeys:")
add_para('     •  Acts 13:4 — "sent on their way by the Holy Spirit"')
add_para('     •  Acts 13:9 — Paul "filled with the Holy Spirit"')
add_para("     •  Acts 16:6, 9 — Later, the Holy Spirit blocks Paul from going to Asia and redirects him to Macedonia through a vision.")
add_para("")
add_rich_para([
    ("Elder Elton said something simple but really important: ", False, False, None),
    ('"God is always calling — are we responding?"', True, True, None),
])
add_para("")
add_para("That has two parts:")
add_para("")
add_rich_para([
    ("(1) Are you actually listening? ", True, False, None),
    ("It’s hard to hear the Holy Spirit if you’re scrolling for 6 hours a day. Putting down \"your 5x2 inch screen\" — for prayer, for the Bible, for silence, for worship — is how you tune your ears to actually hear.", False, False, None),
])
add_para("")
add_rich_para([
    ("(2) When you do hear, do you respond? ", True, False, None),
    ('Sometimes we hear God nudging us in a direction and we just… ignore Him. Like ghosting a text. ("I would never ghost you if you messaged me… na ahh!") Why? Because what He’s asking is uncomfortable. Time-consuming. Costly. But that’s exactly what being a disciple looks like — a "living sacrifice."', False, False, None),
])
add_para("")
add_rich_para([
    ("Bottom line: ", True, False, None),
    ("God already has a plan. You don’t need to invent one. You just need to listen and figure out how to be part of it.", False, True, None),
])

# -- Key Idea 2 --
doc.add_heading("Key Idea 2 — The Church Responds: As SENDER and as SENT", level=2)
add_para("The Antioch church responded to the Holy Spirit in two big ways.")

doc.add_heading("(A) The Church as SENDER", level=3)
add_para("The Antioch church gave to the mission in three ways:")
add_para("•  Their resources — Antioch was a wealthy Roman city. They put their money where their faith was. (ECC parallel: our church is in Bellevue/Redmond/Seattle — also wealthy. ECC’s Missions Fund gives roughly half a million dollars a year to support 40+ long-term missionaries in 17 countries, 7 of them from our own church.)")
add_para("•  Their best people — They didn’t send their leftovers. They sent Barnabas and Paul — the two best teachers in the church. They held nothing back.")
add_para("•  Their support — They prayed, fasted, and encouraged the missionaries on the field. (ECC parallel: every week the bulletin highlights one missionary for prayer. Some small groups have “adopted” a missionary and stay in touch with them.)")

doc.add_heading("(B) The Church as SENT", level=3)
add_para("What kind of people did the Holy Spirit pick? Look at Barnabas and Paul:")
add_rich_para([
    ("•  ", False, False, None),
    ("Barnabas", True, False, None),
    (' — His name literally means "Son of Encouragement." He was the guy who built people up. He’s the one who first gave Paul a chance after Paul’s conversion when nobody else trusted him.', False, False, None),
])
add_rich_para([
    ("•  ", False, False, None),
    ("Paul", True, False, None),
    (' — Plot twist: he used to be a persecutor of the church. Like, the persecution of Stephen that scattered believers to Antioch (last week’s lesson)? He was part of that. But Jesus zapped him on the road to Damascus and called him "a chosen instrument of mine to carry my name before the Gentiles" (Acts 9:15). Now here he is being SENT by the very church he used to hunt down.', False, False, None),
])
add_para("")
add_rich_para([
    ('Both of them shared these traits, and these are the attributes of "', False, False, None),
    ("the Sent", True, False, None),
    ('":', False, False, None),
])
add_para('     \U0001f3af  Missional — willing to share the Gospel with anyone, anywhere')
add_para('     \U0001f4c5  Available — God’s agenda was their agenda')
add_para('     \U0001f504  Flexible / Adaptable — willing to drop their plans when God called')
add_para('     ✋  Willing — they said yes')
add_para('     \U0001f49d  Heart of a servant — selfless. Paul was flogged, beaten with rods, pelted with stones, shipwrecked, jailed multiple times, and left for dead (2 Corinthians 11) — and stayed faithful. He literally said, "For me to live is Christ, and to die is gain" (Philippians 1:21).')
add_para("")
add_rich_para([
    ("Bottom line: ", True, False, None),
    ('Whether God calls you to send or to be sent, He’ll grow these same attributes in you. They aren’t just "missionary traits" — they’re disciple traits.', False, True, None),
])

# -- Key Idea 3 --
doc.add_heading("Key Idea 3 — Your Mission: Learn → Discern → Confirm", level=2)
add_para('Here’s the part that can feel scary or confusing: "Wait, am I supposed to become a missionary in Mongolia or something?"')
add_para("")
add_rich_para([
    ("Elder Elton was super clear: ", False, False, None),
    ("No, not all of us are called to cross-cultural missions. ", True, False, None),
    ("BUT — ", False, False, None),
    ("all of us are called to participate in the Great Commission. ", True, False, None),
    ('Rick Warren said: "The Great Commission is not the Great Suggestion."', False, False, None),
])
add_para("")
add_rich_para([
    ("Start local. ", True, False, None),
    ("Your school, your team, your neighborhood, your friend group, your family — that’s your Jerusalem. The mission field is literally right at your door (and sometimes coming through your door).", False, False, None),
])
add_para("")
add_rich_para([
    ("But don’t stop local. ", True, False, None),
    ('There’s still a huge part of the world (the "10/40 Window") where most people have never heard the Gospel. God may use you in that bigger picture too — eventually.', False, False, None),
])
add_para("")
add_para("So how do you actually figure out your specific mission? Three steps:")

doc.add_heading("1. LEARN — Be a Student", level=3)
add_para('•  Posture of a student: "I’m here to explore."')
add_para("•  Find out what God is already doing in the world.")
add_para("•  Examples: take a class like Perspectives, browse joshuaproject.net, go to a missions conference like Missions Fest Seattle (October) or Urbana (for college students).")

doc.add_heading("2. DISCERN — Listen for the Pattern", level=3)
add_para("Ask yourself five questions:")
add_para("     \U0001f525  Passion — Do I actually care about this? (God rarely makes you do something you hate.)")
add_para("     \U0001f381  Gifting — Do I have the skills / abilities? Does this fit how God made me?")
add_para('     \U0001f6aa  Opportunity — Is God opening a door? (Sometimes "the call" is just someone asking you to help. That’s how Paul recruited Silas and Timothy.)')
add_para("     \U0001f64f  Prayer + Holy Spirit — Ask God to open your eyes. (Warning: He may answer.)")
add_para("     \U0001f474  Wise Counsel — Ask mature Christians what they see in you.")

doc.add_heading("3. CONFIRM — Try It", level=3)
add_para("•  Just go and do. Try the ministry. See if it fits.")
add_para("•  Watch for the Holy Spirit confirming your choice.")
add_para("•  Watch for others affirming your role.")
add_para("•  Pastor Steve Moy was recently ordained — but he was already doing pastoral work for years. Ordination didn’t make him a pastor; it confirmed what God was already doing in him.")
add_para("")
add_rich_para([
    ("Bottom line: ", True, False, None),
    ("You don’t have to wait for a lightning bolt. Most of the time, finding your mission looks like a loop: try, listen, get feedback, adjust, try again.", False, True, None),
])

# -- Small Group Discussion --
doc.add_heading("SMALL GROUP DISCUSSION (30 min)", level=2)
add_para("Split students into groups of 4-5 with one leader.", bold=True)
add_para("Leaders should focus on conversation rather than rushing through questions. Let the good conversations breathe — it’s OK to skip some questions.")
add_para("Encourage students to share honestly and ask questions. No judgment zone.")

# Warm-Up
doc.add_heading("Warm-Up Questions", level=3)
add_para("1. What’s ONE thing you remember from Elder Elton’s sermon? (Even random details count — bonus points if you remember the Mission: Impossible reference, the Antioch church, or what “ghosting God” meant!)")
add_para("")
add_para("2. Be honest: how easy or hard is it for YOU to actually hear God right now? What kinds of things make it harder? What kinds of things make it easier?")
add_para("")
add_para("3. Was there anything from the sermon that confused you, surprised you, or that you want to ask about?")

# Understanding
doc.add_heading("Understanding the Topic", level=3)
add_para("4. The church at Antioch wasn’t having a missions meeting when the Holy Spirit spoke — they were worshiping and fasting. Why does that matter? What does it tell you about how God usually communicates with us? (Hint: it’s not always when we ask.)")
add_para("")
add_rich_para([
    ('5. Elder Elton said: "God is always calling — are we responding?" What’s the difference between ', False, False, None),
    ("listening", True, False, None),
    (" to God and ", False, False, None),
    ("responding", True, False, None),
    (" to Him? Can you do one without the other?", False, False, None),
])
add_para("")
add_para("6. Look at how the Antioch church gave: their resources, their best people, and their prayer/encouragement. Which of those three do you think is hardest for a church (or a person) to actually do? Why?")
add_para("")
add_rich_para([
    ('7. Look at the attributes of "the Sent": missional, available, flexible, willing, servant-hearted. Which of these is ', False, False, None),
    ("most natural", True, False, None),
    (" to you right now? Which is ", False, False, None),
    ("hardest", True, False, None),
    ("? Why?", False, False, None),
])

# Faith vs Real Life
doc.add_heading("Faith vs Real Life", level=3)
add_para('8. Elder Elton talked about "ghosting God" — hearing the Holy Spirit nudge you and just… not responding. Have you ever done this? (Or: do you know what that nudge feels like?) What usually keeps you from saying yes? (Be real — is it fear? Effort? Not wanting to look weird? All of the above?)')
add_para("")
add_para("9. Paul used to be a persecutor — he was literally hunting down Christians. And then God flipped him into the most famous missionary ever. If God can use Paul, what does that say about your past, your weaknesses, or the parts of yourself you wish were different? What does it say about people you might write off?")
add_para("")
add_rich_para([
    ("10. Elder Elton said ", False, False, None),
    ("most", True, False, None),
    (" of us aren’t called to cross-cultural missions overseas — but ", False, False, None),
    ("all", True, False, None),
    (' of us are called to the Great Commission right where we are. What does that look like for you specifically? Is your "mission field" your school? Your team? Your family? Your group chat? Be specific.', False, False, None),
])

# Real Life Application
doc.add_heading("Real Life Application", level=3)
add_para("11. Learn → Discern → Confirm. Pick the step that you think is most relevant to YOUR life right now and explain why:", bold=True)
add_para("     •  LEARN — I haven’t really explored what God’s doing in the world / in my school / in my own life. I’d benefit from being a student.")
add_para("     •  DISCERN — I’m trying to figure out what God might want me to do. I need to look at my passions, gifts, opportunities, and ask wise people.")
add_para("     •  CONFIRM — I have a sense of what I should do, but I haven’t actually tried it yet. I need to step into it and see.")
add_para("")
add_para('     What’s a concrete step you could take this week in that area? (Not "pray more" — be specific. WHO, WHERE, WHEN.)', italic=True)
add_para("")
add_para('12. Pretend a friend at school says, "I don’t believe God talks to people. That’s not real." How would you describe — in your OWN words, not Christian-speak — what it looks like to actually hear from the Holy Spirit? What would you point to from your own experience or from the Bible?')

# Personal Reflection
doc.add_heading("Personal Reflection", level=3)
add_para('13. Think about the five "discern" questions: Passion, Gifting, Opportunity, Prayer, Wise Counsel. Take a minute and answer just two of them silently for yourself:')
add_para("     •  What’s something I genuinely care about? (Passion)")
add_para("     •  What’s something I’m actually decent at — that other people have noticed? (Gifting)")
add_para("")
add_para("     Now put those two together. What’s something God might be wiring you up to do — even in middle school, even before you have a job or a degree? Share if you’re comfortable.", italic=True)
add_para("")
add_para("14. The Antioch church worshiped and fasted before the Holy Spirit spoke. That’s a posture — a way of life. What’s ONE small habit you could start this week to be the kind of person who’s actually paying attention when God speaks? (Examples: 5 min of silence before bed, no phone for the first 10 min of the day, journaling one prayer, reading one Psalm before school…). It has to be small enough that you’ll actually do it.")
add_para("     (Leaders: have each student share theirs out loud as a commitment.)", italic=True)

# -- Wrap-Up --
doc.add_heading("WRAP-UP & PRAYER (2 min)", level=2)
add_para("Ask students to finish this sentence out loud or in their heads:")
add_para("")
p = doc.add_paragraph(style='Normal')
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run('"My mission this week, right where God has placed me, is to __________."')
run.bold = True
run.italic = True

add_para("")
add_para("Close with prayer:", bold=True)
add_para("")
add_para('"God, thank you for this story in Acts 13. Thank you that the Holy Spirit is the one who initiates — we don’t have to figure it all out on our own. Help us be like the church in Antioch: worshiping, fasting, listening, and ready to respond when you speak. Forgive us for the times we’ve ghosted you — when we heard your voice and looked the other way because what you were asking felt uncomfortable or costly. Make us missional. Make us available. Make us flexible and willing and selfless. Help us start where we are — at our schools, on our teams, in our homes, in our group chats — and be the kind of people who carry your name wherever you scatter us. Thank you that the Great Commission is for all of us, not just for missionaries with passports. Show each one of us what our mission looks like this week. We don’t need to know the whole plan. We just need to take the next step. In Jesus’ name, amen."', italic=True)

# -- Save --
out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "ECC Sunday School - Your Mission in the Great Commission (Middle School).docx")
doc.save(out)
print("Saved -> {}".format(out))
